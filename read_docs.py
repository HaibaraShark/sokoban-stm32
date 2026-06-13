import os
base = r"E:\Project\school\pushBit\报告产出"
files = os.listdir(base)
docx_file = [f for f in files if f.endswith(".docx")][0]
xlsx_file = [f for f in files if f.endswith(".xlsx")][0]

print("=== DOCX ===")
print("File:", docx_file)

import docx
doc = docx.Document(os.path.join(base, docx_file))
for p in doc.paragraphs:
    print(repr(p.text))

print("=== TABLES ===")
for i, t in enumerate(doc.tables):
    print(f"Table {i}:")
    for r in t.rows:
        print([c.text for c in r.cells])

print("=== XLSX ===")
print("File:", xlsx_file)

import openpyxl
wb = openpyxl.load_workbook(os.path.join(base, xlsx_file))
print("Sheets:", wb.sheetnames)
for sname in wb.sheetnames:
    ws = wb[sname]
    print(f"Sheet: {sname}, Merged: {list(ws.merged_cells.ranges)}")
    for i, row in enumerate(ws.iter_rows(min_row=1), 1):
        vals = [cell.value for cell in row]
        if any(v is not None for v in vals):
            print(f"Row {i}: {vals}")
